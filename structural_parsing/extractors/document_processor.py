"""
Document Processor Module
Handles text extraction from various document formats.
"""

import os
import logging
from typing import Optional, Dict, Any
import fitz  # PyMuPDF
from docx import Document
import markdown
from bs4 import BeautifulSoup
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing for multiple file formats."""
    
    ALLOWED_EXTENSIONS = {'adoc', 'md', 'dita', 'docx', 'pdf', 'txt'}
    
    def __init__(self):
        """Initialize the document processor."""
        self.supported_formats = {
            '.pdf': self._extract_from_pdf,
            '.docx': self._extract_from_docx,
            '.md': self._extract_from_markdown,
            '.adoc': self._extract_from_asciidoc,
            '.dita': self._extract_from_dita,
            '.txt': self._extract_from_text
        }
    
    def allowed_file(self, filename: str) -> bool:
        """Check if file extension is allowed."""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.ALLOWED_EXTENSIONS
    
    def extract_text(self, filepath: str) -> Optional[str]:
        """
        Extract text from a document file.
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            if not os.path.exists(filepath):
                logger.error(f"File not found: {filepath}")
                return None
            
            file_ext = os.path.splitext(filepath)[1].lower()
            
            if file_ext not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_ext}")
                return None
            
            # Extract text using appropriate method
            text = self.supported_formats[file_ext](filepath)
            
            if text is not None:
                # Clean and normalize the text
                text = self._clean_text(text)
                logger.info(f"Successfully extracted {len(text)} characters from {filepath}")
                return text
            else:
                logger.error(f"Failed to extract text from {filepath}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {filepath}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, filepath: str) -> Optional[str]:
        """Extract text from PDF file."""
        try:
            doc = fitz.open(filepath)
            text = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                try:
                    # Try newer PyMuPDF API first
                    text += page.get_text()
                except AttributeError:
                    # Fallback to older API
                    text += page.getText()
                text += "\n\n"  # Add page break
            
            doc.close()
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from PDF: {str(e)}")
            return None
    
    def _extract_from_docx(self, filepath: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            doc = Document(filepath)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from DOCX: {str(e)}")
            return None
    
    def _extract_from_markdown(self, filepath: str) -> Optional[str]:
        """Extract text from Markdown file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # Convert markdown to HTML, then extract text
            html = markdown.markdown(md_content)
            soup = BeautifulSoup(html, 'html.parser')
            text = soup.get_text()
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from Markdown: {str(e)}")
            return None
    
    def _extract_from_asciidoc(self, filepath: str) -> Optional[str]:
        """Extract text from AsciiDoc file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Remove AsciiDoc formatting (basic implementation)
            # Convert headers to text (preserve content, remove =+ markers)
            content = re.sub(r'^=+\s+(.*?)$', r'\1', content, flags=re.MULTILINE)
            
            # Remove attribute entries (lines starting with :attribute:)
            content = re.sub(r'^:[\w-]+:\s*.*?$', '', content, flags=re.MULTILINE)
            
            # Remove block delimiters but preserve content inside
            content = re.sub(r'^----+$', '', content, flags=re.MULTILINE)
            content = re.sub(r'^\*\*\*\*+$', '', content, flags=re.MULTILINE)
            content = re.sub(r'^====+$', '', content, flags=re.MULTILINE)
            content = re.sub(r'^\+\+\+\++$', '', content, flags=re.MULTILINE)
            
            # Remove admonition markers but preserve content
            content = re.sub(r'^\[(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]$', '', content, flags=re.MULTILINE)
            
            # Remove inline formatting
            content = re.sub(r'\*([^*]+)\*', r'\1', content)  # bold
            content = re.sub(r'_([^_]+)_', r'\1', content)    # italic
            content = re.sub(r'`([^`]+)`', r'\1', content)    # monospace
            
            # Remove links but keep text
            content = re.sub(r'link:([^[]+)\[([^\]]*)\]', r'\2', content)
            content = re.sub(r'http[s]?://[^\s\[\]]+', '', content)
            
            # Clean up extra whitespace
            content = re.sub(r'\n\s*\n', '\n\n', content)
            content = re.sub(r'\n{3,}', '\n\n', content)
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from AsciiDoc: {str(e)}")
            return None
    
    def _extract_from_dita(self, filepath: str) -> Optional[str]:
        """Extract text from DITA file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Parse DITA XML and extract text
            soup = BeautifulSoup(content, 'xml')
            
            # Remove unwanted metadata but preserve titles
            for element in soup(['prolog', 'metadata']):
                element.decompose()
            
            # Extract text content
            text = soup.get_text(separator='\n')
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from DITA: {str(e)}")
            return None
    
    def _extract_from_text(self, filepath: str) -> Optional[str]:
        """Extract text from plain text file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Error extracting from text file: {str(e)}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if text is None:
            return ""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Trim whitespace
        text = text.strip()
        
        return text
    
    def get_document_info(self, filepath: str) -> Dict[str, Any]:
        """
        Get document information and metadata.
        
        Args:
            filepath: Path to the document file
            
        Returns:
            Dictionary with document information
        """
        info = {
            'filepath': filepath,
            'filename': os.path.basename(filepath),
            'file_size': 0,
            'format': 'unknown',
            'extractable': False
        }
        
        try:
            # Extract format from file extension regardless of file existence
            file_ext = os.path.splitext(filepath)[1].lower()
            if file_ext:
                info['format'] = file_ext.lstrip('.')
                info['extractable'] = file_ext in self.supported_formats
            
            if os.path.exists(filepath):
                info['file_size'] = os.path.getsize(filepath)
            
        except Exception as e:
            logger.error(f"Error getting document info: {str(e)}")
        
        return info 